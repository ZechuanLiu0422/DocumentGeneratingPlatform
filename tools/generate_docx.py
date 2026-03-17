#!/usr/bin/env python3
"""
Word 文档生成引擎
"""
import json
import sys
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from format_utils import format_chinese_date

def setup_page(doc, page_setup):
    """设置页面格式"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.top_margin = Cm(page_setup.get('top_margin', 2.54))
    section.bottom_margin = Cm(page_setup.get('bottom_margin', 2.54))
    section.left_margin = Cm(page_setup.get('left_margin', 3.17))
    section.right_margin = Cm(page_setup.get('right_margin', 3.17))

def apply_style(paragraph, style_config):
    """应用样式到段落"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

    # 字体
    font_name = style_config.get('font', '仿宋_GB2312')
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 字号
    run.font.size = Pt(style_config.get('size', 16))

    # 加粗
    if style_config.get('bold', False):
        run.font.bold = True

    # 行间距（支持磅值或倍数）
    line_spacing = style_config.get('line_spacing', 1.5)
    line_spacing_rule = style_config.get('line_spacing_rule', 'multiple')
    if line_spacing_rule == 'exactly':
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    else:
        paragraph.paragraph_format.line_spacing = line_spacing

def detect_heading_level(line):
    """检测行是否为标题，返回标题级别或None"""
    import re
    line = line.strip()

    # 一级标题：一、二、三、四、五、六、七、八、九、十
    if re.match(r'^[一二三四五六七八九十]+、', line):
        return 'heading1'

    # 二级标题：（一）（二）（三）
    if re.match(r'^[（(][一二三四五六七八九十]+[）)]', line):
        return 'heading2'

    return None

def render_attachments(doc, attachments, styles):
    """渲染附件列表"""
    style_config = styles.get('body', {})

    # 第一行："附件：" + 第一个附件
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(2 * 16)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)

    run = p.add_run('附件：')
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(16)

    # 第一个附件紧接着
    if attachments:
        run = p.add_run(f"1. {attachments[0]}")
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)

    # 后续附件另起一行并缩进
    for i, att_name in enumerate(attachments[1:], start=2):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(5 * 16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)

        run = p.add_run(f"{i}. {att_name}")
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)

    # 附件列表后添加两个空行
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)
        run = p.add_run(' ')
        run.font.name = '仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        run.font.size = Pt(16)

def render_section(doc, section, content_data, styles):
    """渲染文档的一个部分"""
    field = section['field']

    # 特殊处理附件字段
    if field == 'attachment':
        attachments = content_data.get('attachments', [])
        if attachments and len(attachments) > 0:
            render_attachments(doc, attachments, styles)
        return

    # 特殊处理联系人字段
    if field == 'contact':
        contact_name = content_data.get('contactName', '')
        contact_phone = content_data.get('contactPhone', '')
        if contact_name and contact_phone:
            content = f"（联系人：{contact_name}，电话：{contact_phone}）"
        else:
            return  # 如果任一字段为空，不渲染此部分
    else:
        content = content_data.get(field, '')

    # 空行字段：即使内容为空也渲染
    if field.startswith('_blank'):
        paragraph = doc.add_paragraph()
        # 应用body样式到空行
        style_config = styles.get('body', {})
        run = paragraph.add_run(' ')
        font_name = style_config.get('font', '仿宋_GB2312')
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(style_config.get('size', 16))
        # 行间距
        line_spacing = style_config.get('line_spacing', 28)
        line_spacing_rule = style_config.get('line_spacing_rule', 'exactly')
        if line_spacing_rule == 'exactly':
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(line_spacing)
        else:
            paragraph.paragraph_format.line_spacing = line_spacing
        # 段前段后间距强制设置为0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return

    if not content:
        return

    # 内容清理
    import re
    if field == 'content':
        # 正文：保留段落分隔，清理单行内的多余空格
        content = re.sub(r'[ \t]+', ' ', content)  # 清理空格和制表符
        content = re.sub(r'\n[ \t]+', '\n', content)  # 清理行首空格
        content = content.strip()
    else:
        # 其他字段：清理所有多余空白
        content = re.sub(r'\s+', ' ', content).strip()

    # 主送机关自动添加冒号
    if field == 'recipient' and content and not content.endswith('：') and not content.endswith(':'):
        content = content + '：'

    # 日期格式化
    if field == 'date' and content:
        try:
            content = format_chinese_date(content)
        except:
            pass

    # 获取样式配置
    style_name = section.get('style', 'body')
    style_config = styles.get(style_name, {})

    # 按换行符分割，每行创建新段落（硬回车）
    lines = content.split('\n')

    for i, line in enumerate(lines):
        if not line.strip() and i > 0:
            continue

        paragraph = doc.add_paragraph()

        # 检测标题级别
        heading_level = detect_heading_level(line)
        current_style_config = style_config
        if heading_level:
            current_style_config = styles.get(heading_level, style_config)

        # 对齐方式
        align = section.get('align', 'justify')
        if heading_level:
            # 标题左对齐
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 缩进（字符数）
        indent = section.get('indent', 0)
        if indent > 0:
            font_size = current_style_config.get('size', 16)
            paragraph.paragraph_format.first_line_indent = Pt(indent * font_size)

        # 右空字符处理
        right_space_chars = section.get('right_space_chars', 0)
        if right_space_chars > 0:
            paragraph.paragraph_format.right_indent = Cm(right_space_chars * 0.56)

        # 段前段后间距强制设置为0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # 行间距
        line_spacing = current_style_config.get('line_spacing', 1.5)
        line_spacing_rule = current_style_config.get('line_spacing_rule', 'multiple')
        if line_spacing_rule == 'exactly':
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(line_spacing)
        else:
            paragraph.paragraph_format.line_spacing = line_spacing

        # 添加内容并应用字体样式
        run = paragraph.add_run(line)
        font_name = current_style_config.get('font', '仿宋_GB2312')
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(current_style_config.get('size', 16))
        if current_style_config.get('bold', False):
            run.font.bold = True

def generate_document(template_path, content_data, output_path):
    """生成 Word 文档"""
    # 加载模板配置
    with open(template_path, 'r', encoding='utf-8') as f:
        template = json.load(f)

    # 创建文档
    doc = Document()

    # 设置页面
    setup_page(doc, template['page_setup'])

    # 渲染各个部分
    for section in template['sections']:
        render_section(doc, section, content_data, template['styles'])

    # 保存
    doc.save(output_path)

def main():
    parser = argparse.ArgumentParser(description='Word 文档生成工具')
    parser.add_argument('--template', required=True, help='模板文件路径')
    parser.add_argument('--content', required=True, help='内容数据 (JSON)')
    parser.add_argument('--output', required=True, help='输出文件路径')

    args = parser.parse_args()

    try:
        content_data = json.loads(args.content)
        generate_document(args.template, content_data, args.output)

        result = {
            'success': True,
            'file_path': args.output
        }
        print(json.dumps(result, ensure_ascii=False))

    except Exception as e:
        result = {
            'success': False,
            'error': str(e)
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)

if __name__ == '__main__':
    main()
